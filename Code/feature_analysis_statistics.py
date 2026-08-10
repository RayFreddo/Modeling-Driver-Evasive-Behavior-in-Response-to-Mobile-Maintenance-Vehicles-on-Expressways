from __future__ import annotations

import argparse
import os
from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from scipy.stats import kruskal


ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = ROOT / "数据处理"
MODEL_ROOT = ROOT / "多模型预测输出"
FORMAL_DIR = MODEL_ROOT / "00_结果数据" / "正式完整轨迹结果"
STATS_XLSX = MODEL_ROOT / "00_结果数据" / "重跑统计结果.xlsx"
SOURCE_DOCX = ROOT / "TRB-submission.docx"
OUTPUT_DOCX = ROOT / "TRB-submission_Results_Revised.docx"
TEMP_DOCX = ROOT / ".TRB-submission_Results_Revised.tmp.docx"
TEMP_MEDIA_DOCX = ROOT / ".TRB-submission_Results_Revised.media.tmp.docx"
OUT_DIR = ROOT / "figure" / "results_4k_sg_comment_draft"
SOURCE_DIR = OUT_DIR / "source_data"
ORIGINAL_FIG10 = ROOT / "figure" / "Figure8_Representative_Observed_and_Modeled_Evasive_Trajectories_4K.png"
ORIGINAL_FIG11 = ROOT / "figure" / "Figure9_original_content_upscaled_600dpi.png"
ORIGINAL_FIG12 = ROOT / "figure" / "Figure_time_horizon_error_evolution_vertical.png"
ORIGINAL_FDE = ROOT / "figure" / "analysis" / "多模型预测输出" / "01_整条轨迹指标图" / "Figure_FDE_box.png"
ANNOTATED_FIG10 = OUT_DIR / "Figure10_Original_With_Panel_Titles_4K.png"
ANNOTATED_FIG12 = OUT_DIR / "Figure12_Original_With_Panel_Titles_4K.png"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SOURCE_DIR.mkdir(parents=True, exist_ok=True)

TRAIN_DAYS = ["0403", "0407", "0410", "0416"]
SCENE_ORDER_ZH = [
    "相邻侧无车",
    "相邻侧前方有车后方无车",
    "相邻侧后方有车前方无车",
    "相邻侧前后均有车",
]
SCENE_ORDER_EN = [
    "No vehicle",
    "Leading only",
    "Following only",
    "Both",
]
SCENE_TICK_LABELS = ["No\nvehicle", "Leading\nonly", "Following\nonly", "Both"]
SCENE_ROW_EN = [
    "No adjacent-lane vehicle",
    "Adjacent-lane leading vehicle only",
    "Adjacent-lane following vehicle only",
    "Both adjacent-lane leading and following vehicles",
]
SCENE_COLORS = ["#E7EFF5", "#BED1E0", "#7FA3BF", "#456F91"]

MODEL_ORDER = ["HGB-ECR", "Kalman-CV", "GBR", "IDM", "CNN", "Helly"]
MODEL_LABELS = {
    "HGB-ECR": "HGB-ECR",
    "Kalman-CV": "Kalman-CV",
    "GBR": "GBR",
    "IDM": "IDM-Avoidance",
    "CNN": "1D-CNN",
    "Helly": "Helly Baseline",
}
MODEL_COLORS = {
    "HGB-ECR": "#C83D34",
    "Kalman-CV": "#4C78A8",
    "GBR": "#F28E2B",
    "IDM": "#B07AA1",
    "CNN": "#76B7B2",
    "Helly": "#59A14F",
}
MODEL_LINESTYLES = {
    "HGB-ECR": "-", "Kalman-CV": "-", "GBR": "--",
    "IDM": "-.", "CNN": ":", "Helly": (0, (5, 2)),
}
MODEL_MARKERS = {
    "HGB-ECR": "o", "Kalman-CV": "s", "GBR": "^",
    "IDM": "D", "CNN": "v", "Helly": "P",
}
ACTION_ORDER_ZH = ["换道", "减速后换道", "同时减速换道"]
ACTION_LABELS = ["Direct lane change", "Deceleration then lane change", "Synchronized response"]
ACTION_COLORS = ["#BED1E0", "#587F9F", "#A94442"]


def apply_style():
    plt.rcParams["font.family"] = "sans-serif"
    plt.rcParams["font.sans-serif"] = ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"]
    plt.rcParams["svg.fonttype"] = "none"
    plt.rcParams["pdf.fonttype"] = 42
    # Results figures are reduced to roughly 74-78% of their source width in
    # the TRB layout. These source sizes therefore render close to 10 pt body
    # text after placement in Word.
    plt.rcParams["font.size"] = 11.5
    plt.rcParams["axes.titlesize"] = 12.0
    plt.rcParams["axes.labelsize"] = 12.0
    plt.rcParams["xtick.labelsize"] = 10.5
    plt.rcParams["ytick.labelsize"] = 10.5
    plt.rcParams["legend.fontsize"] = 10.0
    plt.rcParams["axes.spines.right"] = False
    plt.rcParams["axes.spines.top"] = False
    plt.rcParams["axes.linewidth"] = 0.8
    plt.rcParams["legend.frameon"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"


def save_figure(fig, stem: str, dpi: int = 600):
    base = OUT_DIR / stem
    fig.savefig(base.with_suffix(".png"), dpi=dpi, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return base.with_suffix(".png")


def panel_label(ax, label, title=None):
    text_value = label if title is None else f"{label}  {title}"
    ax.text(
        0.012, 0.988, text_value, transform=ax.transAxes,
        ha="left", va="top", fontsize=11.5 if title else 12.5,
        fontweight="bold", zorder=20,
        bbox={"facecolor": "white", "edgecolor": "none", "alpha": 0.86, "pad": 1.2},
    )


def annotate_original_figures():
    """Add panel titles only in whitespace of copies of the original figures."""
    regular = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
    bold = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"

    image10 = Image.open(ORIGINAL_FIG10).convert("RGB")
    draw10 = ImageDraw.Draw(image10)
    font10 = ImageFont.truetype(bold, 42)
    draw10.text((165, 82), "Representative lane-changing trajectory", fill="black", font=font10)
    draw10.text((2175, 82), "Representative continuous lateral adjustment", fill="black", font=font10)
    image10.save(ANNOTATED_FIG10, dpi=(600, 600))

    image12 = Image.open(ORIGINAL_FIG12).convert("RGB")
    draw12 = ImageDraw.Draw(image12)
    font12 = ImageFont.truetype(bold, 52)
    labels12 = [
        ((285, 310), "a  ADE over prediction horizon"),
        ((285, 1800), "b  FDE over prediction horizon"),
        ((285, 3315), "c  RMSE over prediction horizon"),
    ]
    for (x, y), text_value in labels12:
        bbox = draw12.textbbox((x, y), text_value, font=font12)
        draw12.rectangle((bbox[0] - 10, bbox[1] - 5, bbox[2] + 10, bbox[3] + 5), fill="white")
        draw12.text((x, y), text_value, fill="black", font=font12)
    image12.save(ANNOTATED_FIG12, dpi=(600, 600))
    return ANNOTATED_FIG10, ANNOTATED_FIG12


def style_axis(ax, grid_axis="y"):
    ax.grid(axis=grid_axis, color="#D9D9D9", linewidth=0.6, alpha=0.65)
    ax.set_axisbelow(True)


def colored_boxplot(ax, groups, labels, colors, ylabel, ylim=None):
    bp = ax.boxplot(
        groups,
        tick_labels=labels,
        patch_artist=True,
        showfliers=False,
        whis=(5, 95),
        medianprops={"color": "#272727", "linewidth": 1.4},
        whiskerprops={"color": "#606060", "linewidth": 0.9},
        capprops={"color": "#606060", "linewidth": 0.9},
        boxprops={"edgecolor": "#606060", "linewidth": 0.9},
    )
    for patch, color in zip(bp["boxes"], colors):
        patch.set_facecolor(color)
        patch.set_alpha(0.9)
    ax.set_ylabel(ylabel)
    ax.tick_params(axis="x", rotation=0)
    if ylim is not None:
        ax.set_ylim(*ylim)
    style_axis(ax)
    return bp


def read_training_data():
    acceleration_parts = []
    trajectory_parts = []
    record_counts = {}
    usecols = [
        "object_id", "segment_id", "timestamp_s", "distance_m", "follower_speed_mps",
        "follower_acc_mps2", "ttc_s", "x_m", "y_m",
    ]
    for day in TRAIN_DAYS:
        path = DATA_DIR / f"{day}_断点插值补全.xlsx"
        df = pd.read_excel(path, sheet_name="Sheet1", usecols=usecols)
        for col in usecols:
            if col not in {"object_id", "segment_id"}:
                df[col] = pd.to_numeric(df[col], errors="coerce")
        df = df.sort_values(["object_id", "segment_id", "timestamp_s"])
        record_counts[day] = int(len(df))
        acceleration_parts.append(df["follower_acc_mps2"].dropna().to_numpy(float))
        grouped = df.groupby(["object_id", "segment_id"], sort=False)
        first = grouped.first().reset_index()
        last = grouped.last().reset_index()
        merged = first.merge(last, on=["object_id", "segment_id"], suffixes=("_start", "_end"))
        merged["day"] = day
        merged["sample_key"] = (
            day + "_" + merged["object_id"].astype(str) + "__seg" + merged["segment_id"].astype(str)
        )
        merged["spacing_change_m"] = merged["distance_m_end"] - merged["distance_m_start"]
        merged["speed_change_mps"] = merged["follower_speed_mps_end"] - merged["follower_speed_mps_start"]
        trajectory_parts.append(merged)

    accelerations = np.concatenate(acceleration_parts)
    trajectory_change = pd.concat(trajectory_parts, ignore_index=True)
    trajectory_change.to_csv(SOURCE_DIR / "training_trajectory_changes.csv", index=False)

    detail = pd.read_excel(STATS_XLSX, sheet_name="五天动作与场景明细")
    detail = detail[detail["day"].isin([403, 407, 410, 416])].copy()
    detail["scene_type"] = pd.Categorical(detail["scene_type"], SCENE_ORDER_ZH, ordered=True)
    detail = detail.sort_values("scene_type").reset_index(drop=True)
    detail.to_csv(SOURCE_DIR / "training_evasive_event_indicators.csv", index=False)

    bins = np.linspace(-5, 3, 161)
    hist, edges = np.histogram(accelerations[np.isfinite(accelerations)], bins=bins)
    pd.DataFrame({"bin_left": edges[:-1], "bin_right": edges[1:], "count": hist}).to_csv(
        SOURCE_DIR / "longitudinal_acceleration_histogram.csv", index=False
    )
    return accelerations, trajectory_change, detail, record_counts


def scene_groups(detail, column):
    return [detail.loc[detail["scene_type"].eq(scene), column].dropna().to_numpy(float) for scene in SCENE_ORDER_ZH]


def scene_summary_and_tests(detail):
    metrics = [
        ("start_gap_m", "Spacing at evasive-event onset"),
        ("start_speed_mps", "Speed at evasive-event onset"),
        ("start_ttc_s", "TTC at evasive-event onset"),
        ("decel_start_distance_m", "Spacing at deceleration initiation"),
        ("lane_start_distance_m", "Spacing at lane-change initiation"),
        ("lane_duration_s", "Lane-change duration"),
        ("speed_drop_mps", "Speed reduction"),
    ]
    rows = []
    for scene_zh, scene_en in zip(SCENE_ORDER_ZH, SCENE_ROW_EN):
        g = detail[detail["scene_type"].eq(scene_zh)]
        rows.append({
            "Moving bottleneck scenario": scene_en,
            "Evasive events, n": len(g),
            "Spacing at evasive-event onset (m)": g["start_gap_m"].median(),
            "Speed at evasive-event onset (m/s)": g["start_speed_mps"].median(),
            "TTC at evasive-event onset (s)": g["start_ttc_s"].median(),
            "Spacing at deceleration initiation (m)": g["decel_start_distance_m"].median(),
            "Spacing at lane-change initiation (m)": g["lane_start_distance_m"].median(),
            "Lane-change duration (s)": g["lane_duration_s"].median(),
            "Speed reduction (m/s)": g["speed_drop_mps"].median(),
        })
    summary = pd.DataFrame(rows)
    summary.to_csv(SOURCE_DIR / "table7_scenario_medians.csv", index=False)

    test_rows = []
    for column, label in metrics:
        groups = scene_groups(detail, column)
        h, p = kruskal(*groups)
        test_rows.append({"Behavioral indicator": label, "H statistic": h, "p value": p})
    tests = pd.DataFrame(test_rows)
    tests.to_csv(SOURCE_DIR / "table8_kruskal_wallis.csv", index=False)
    return summary, tests


def make_figure8(accelerations, detail, record_counts):
    fig, axes = plt.subplots(2, 2, figsize=(8.0, 5.6), constrained_layout=True)
    ax = axes[0, 0]
    vals = accelerations[np.isfinite(accelerations)]
    # The -5 m/s² mass is a lower-bound clipping artefact rather than a
    # measured distributional mode. Exclude that exact clipped value from the
    # displayed histogram while leaving all source data unchanged.
    vals = vals[(vals >= -4.0) & (vals <= 3.0)]
    mean_acc = float(np.mean(vals))
    median_acc = float(np.median(vals))
    ax.hist(vals, bins=72, density=True, color="#9DBAD0", edgecolor="white", linewidth=0.2)
    ax.axvline(mean_acc, color="#A94442", linestyle="--", linewidth=1.5,
               label=f"Mean = {mean_acc:.2f}")
    ax.axvline(median_acc, color="#244F75", linestyle="-", linewidth=1.5,
               label=f"Median = {median_acc:.2f}")
    ax.set_xlabel("Acceleration (m/s²)")
    ax.set_ylabel("Density")
    # Keep the statistics legend below the long panel title. The opaque white
    # box also prevents the histogram bars from showing through the labels.
    ax.legend(
        loc="upper right", bbox_to_anchor=(0.98, 0.82), borderaxespad=0.0,
        fontsize=10.0, frameon=True, framealpha=0.88,
        facecolor="white", edgecolor="none",
    )
    style_axis(ax)
    panel_label(ax, "a", "Acceleration distribution")

    ax = axes[0, 1]
    gap = pd.to_numeric(detail["start_gap_m"], errors="coerce").to_numpy(float)
    speed = pd.to_numeric(detail["start_speed_mps"], errors="coerce").to_numpy(float)
    finite = np.isfinite(gap) & np.isfinite(speed)
    gap, speed = gap[finite], speed[finite]
    gap_lo, gap_hi = np.nanpercentile(gap, [2, 98])
    speed_lo, speed_hi = np.nanpercentile(speed, [2, 98])
    shown = (gap >= gap_lo) & (gap <= gap_hi) & (speed >= speed_lo) & (speed <= speed_hi)
    hb = ax.hexbin(gap[shown], speed[shown], gridsize=34,
                   mincnt=1, cmap="Blues", bins="log", linewidths=0)
    cbar = fig.colorbar(hb, ax=ax, fraction=0.036, pad=0.02, aspect=28)
    cbar.set_label("Count")
    ax.set_xlim(gap_lo, gap_hi)
    ax.set_ylim(speed_lo, speed_hi)
    ax.set_xlabel("Spacing (m)")
    ax.set_ylabel("Speed (m/s)")
    style_axis(ax, grid_axis="both")
    panel_label(ax, "b", "Spacing-speed distribution")

    ax = axes[1, 0]
    groups = scene_groups(detail, "start_ttc_s")
    ymax = float(np.nanpercentile(np.concatenate(groups), 95) * 1.12)
    colored_boxplot(ax, groups, SCENE_TICK_LABELS, SCENE_COLORS, "TTC (s)", (0, ymax))
    panel_label(ax, "c", "TTC by scenario")

    ax = axes[1, 1]
    groups = scene_groups(detail, "decel_start_distance_m")
    ymax = float(np.nanpercentile(np.concatenate(groups), 95) * 1.10)
    colored_boxplot(ax, groups, SCENE_TICK_LABELS, SCENE_COLORS, "Distance (m)", (0, ymax))
    panel_label(ax, "d", "Deceleration initiation distance")
    return save_figure(fig, "Figure8_Longitudinal_Risk_Indicators_4K")


def make_figure9(detail, trajectory_change):
    fig, axes = plt.subplots(3, 2, figsize=(9.5, 7.8), constrained_layout=True)
    specs = [
        ("lane_start_distance_m", "Distance (m)"),
        ("lane_duration_s", "Duration (s)"),
        ("lat_shift_m", "Displacement (m)"),
        ("min_gap_m", "Gap (m)"),
    ]
    panel_titles = [
        "Lane-change initiation distance",
        "Lane-change duration",
        "Lateral displacement",
        "Minimum adjacent-lane gap",
    ]
    for label, title, ax, (column, ylabel) in zip("abcd", panel_titles, axes.flat[:4], specs):
        groups = scene_groups(detail, column)
        groups = [np.abs(g) if column == "lat_shift_m" else g for g in groups]
        ymax = float(np.nanpercentile(np.concatenate(groups), 95) * 1.12)
        colored_boxplot(ax, groups, SCENE_TICK_LABELS, SCENE_COLORS, ylabel, (0, ymax))
        panel_label(ax, label, title)

    ax = axes[2, 0]
    x = pd.to_numeric(trajectory_change["spacing_change_m"], errors="coerce")
    y = pd.to_numeric(trajectory_change["speed_change_mps"], errors="coerce")
    mask = x.notna() & y.notna()
    xv, yv = x[mask].to_numpy(float), y[mask].to_numpy(float)
    xlo, xhi = np.nanpercentile(xv, [1, 99])
    ylo, yhi = np.nanpercentile(yv, [1, 99])
    shown = (xv >= xlo) & (xv <= xhi) & (yv >= ylo) & (yv <= yhi)
    hb = ax.hexbin(xv[shown], yv[shown], gridsize=38, mincnt=1, cmap="Blues", bins="log", linewidths=0)
    ax.axhline(0, color="#606060", linewidth=0.8)
    ax.axvline(0, color="#606060", linewidth=0.8)
    corr = float(np.corrcoef(xv, yv)[0, 1])
    coupled = float(np.mean((xv < 0) & (yv < 0)) * 100)
    ax.text(0.97, 0.06, f"Pearson r = {corr:.2f}\nBoth decreased: {coupled:.1f}%",
            transform=ax.transAxes, va="bottom", ha="right", fontsize=9.8,
            bbox={"facecolor": "white", "edgecolor": "none", "alpha": 0.82, "pad": 1.2})
    ax.set_xlabel("Spacing change (m)")
    ax.set_ylabel("Speed change (m/s)")
    style_axis(ax, grid_axis="both")
    panel_label(ax, "e", "Spacing-speed coupling")

    ax = axes[2, 1]
    composition = pd.crosstab(detail["scene_type"], detail["action"], normalize="index").reindex(
        index=SCENE_ORDER_ZH, columns=ACTION_ORDER_ZH, fill_value=0
    ) * 100
    bottom = np.zeros(len(composition))
    xloc = np.arange(len(composition))
    for action, action_label, color in zip(ACTION_ORDER_ZH, ACTION_LABELS, ACTION_COLORS):
        vals = composition[action].to_numpy(float)
        bars = ax.bar(xloc, vals, bottom=bottom, color=color, edgecolor="white", linewidth=0.6, label=action_label)
        for bar, value, base in zip(bars, vals, bottom):
            if value >= 9:
                ax.text(bar.get_x() + bar.get_width() / 2, base + value / 2, f"{value:.1f}%",
                        ha="center", va="center", fontsize=9.5,
                        color="white" if color == "#B64342" else "#272727")
        bottom += vals
    ax.set_xticks(xloc, SCENE_TICK_LABELS)
    ax.set_ylim(0, 100)
    ax.set_ylabel("Proportion (%)")
    handles, labels = ax.get_legend_handles_labels()
    ax.legend(handles, labels, loc="center left", bbox_to_anchor=(1.02, 0.50),
              ncol=1, fontsize=9.5, handlelength=1.5, columnspacing=1.3,
              borderaxespad=0.0, handletextpad=0.5)
    style_axis(ax)
    ax.set_title("f  Evasive-strategy composition", loc="left", pad=7,
                 fontsize=11.5, fontweight="bold")
    composition.rename(index=dict(zip(SCENE_ORDER_ZH, SCENE_ROW_EN)), columns=dict(zip(ACTION_ORDER_ZH, ACTION_LABELS))).to_csv(
        SOURCE_DIR / "evasive_strategy_composition.csv"
    )
    pd.DataFrame({"spacing_change_m": xv, "speed_change_mps": yv}).to_csv(
        SOURCE_DIR / "spacing_speed_coupling.csv", index=False
    )
    return save_figure(fig, "Figure9_Lateral_Surrounding_Strategy_1200dpi_draft", dpi=1200)


def load_model_results():
    summary = pd.read_csv(FORMAL_DIR / "六模型完整轨迹正式指标.csv")
    summary = summary[summary["model_id"].isin(MODEL_ORDER)].copy()
    summary["model_label_en"] = summary["model_id"].map(MODEL_LABELS)
    summary["order"] = summary["model_id"].map({m: i for i, m in enumerate(MODEL_ORDER)})
    summary = summary.sort_values("order")

    time_data = pd.read_csv(FORMAL_DIR / "各时间节点ADE_FDE_RMSE.csv")
    time_data = time_data[time_data["model_id"].isin(MODEL_ORDER)].copy()
    time_data["model_label_en"] = time_data["model_id"].map(MODEL_LABELS)

    workbook = FORMAL_DIR / "六模型整条轨迹累计预测正式结果.xlsx"
    point = pd.read_excel(workbook, sheet_name="逐时刻误差")
    point = point[point["model_id"].isin(MODEL_ORDER)].copy()
    object_metrics = pd.read_excel(workbook, sheet_name="逐车整条轨迹指标")
    object_metrics = object_metrics[object_metrics["model_id"].isin(MODEL_ORDER)].copy()
    window = pd.read_excel(STATS_XLSX, sheet_name="时间窗对比")

    summary.to_csv(SOURCE_DIR / "model_performance_summary.csv", index=False)
    time_data.to_csv(SOURCE_DIR / "model_error_by_horizon.csv", index=False)
    window.to_csv(SOURCE_DIR / "history_window_sensitivity.csv", index=False)
    return summary, time_data, point, object_metrics, window


def trajectory_panel(ax, point, sample_key, panel, descriptor):
    true = point[
        (point["sample_key"].eq(sample_key)) & (point["model_id"].eq("HGB-ECR"))
    ].sort_values("rel_time_s")
    true_x = true["true_x_m"].to_numpy(float)
    true_y = true["true_y_m"].to_numpy(float)
    reference_x, reference_y = true_x[0], true_y[0]
    ax.plot(true_x, true_y, color="#202020", linewidth=2.6,
            label="Observed", zorder=8)
    figure_labels = {**MODEL_LABELS, "Helly": "Helly"}
    for model in MODEL_ORDER:
        g = point[(point["sample_key"].eq(sample_key)) & (point["model_id"].eq(model))].sort_values("rel_time_s")
        if g.empty:
            continue
        pred_x = g["pred_x_m"].to_numpy(float)
        pred_y = g["pred_y_m"].to_numpy(float)
        # Translate displayed predictions to the observed initial point for
        # visual comparison; source coordinates and ADE/FDE remain unchanged.
        pred_x = pred_x + (reference_x - pred_x[0])
        pred_y = pred_y + (reference_y - pred_y[0])
        ax.plot(
            pred_x, pred_y, color=MODEL_COLORS[model],
            linewidth=2.35 if model == "HGB-ECR" else 1.45,
            alpha=1.0 if model == "HGB-ECR" else 0.92,
            label=figure_labels[model],
        )
    ade = float(true["error_m"].mean())
    fde = float(true.iloc[-1]["error_m"])
    sample_label = sample_key.replace("__seg", "-seg")
    ax.set_xlabel("Aligned lateral position x / m")
    ax.set_ylabel("Aligned longitudinal position y / m")
    ax.set_title(
        f"{sample_label}  |  HGB-ECR ADE={ade:.2f} m, FDE={fde:.2f} m",
        pad=6, fontsize=10.5,
    )
    ax.text(0.0, 1.09, f"({panel})  {descriptor}", transform=ax.transAxes,
            ha="left", va="bottom", fontsize=11.0, fontweight="bold")
    style_axis(ax, grid_axis="both")


def make_figure10(point, window):
    del window
    fig, axes = plt.subplots(1, 2, figsize=(8.0, 3.8))
    trajectory_panel(
        axes[0], point, "1629__seg1", "a", "Representative lane-changing trajectory"
    )
    trajectory_panel(
        axes[1], point, "11535__seg1", "b", "Representative continuous lateral adjustment"
    )
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(
        handles, labels, loc="lower center", bbox_to_anchor=(0.5, 0.015),
        ncol=4, fontsize=9.2, handlelength=2.0, columnspacing=1.0,
    )
    fig.subplots_adjust(left=0.08, right=0.98, bottom=0.24, top=0.80, wspace=0.28)
    return save_figure(fig, "Figure10_Original_Layout_Larger_Fonts_4K")


def make_figure11(summary):
    data = summary.set_index("model_id").loc[MODEL_ORDER].reset_index()
    with plt.rc_context({
        "font.size": 16.5,
        "axes.titlesize": 18.0,
        "axes.labelsize": 17.0,
        "xtick.labelsize": 15.5,
        "ytick.labelsize": 15.5,
        "legend.fontsize": 14.5,
    }):
        fig, ax = plt.subplots(figsize=(8.0, 4.65), constrained_layout=True)
        x = np.arange(len(data))
        width = 0.23
        colors = ["#D7E5F0", "#91B8D8", "#4F84B5"]
        for offset, column, label, color in zip(
            [-1, 0, 1], ["ADE_m", "FDE_m", "RMSE_m"], ["ADE", "FDE", "RMSE"], colors
        ):
            ax.bar(x + offset * width, data[column], width=width, color=color,
                   edgecolor="#4F84B5", linewidth=0.7, label=label)
        ax.set_xticks(x, data["model_label_en"], rotation=22, ha="right")
        ax.set_ylabel("Error / m")
        ax.set_title("Overall Trajectory Prediction Error")
        ax.legend(loc="upper left", ncol=3)
        style_axis(ax)
    return save_figure(fig, "Figure11_Original_Layout_Larger_Fonts_4K")


def make_figure12(time_data):
    with plt.rc_context({
        "font.size": 13.0,
        "axes.titlesize": 13.5,
        "axes.labelsize": 13.5,
        "xtick.labelsize": 12.0,
        "ytick.labelsize": 12.0,
        "legend.fontsize": 11.5,
    }):
        fig, axes = plt.subplots(3, 1, figsize=(8.0, 9.4), sharex=True)
        specs = [
            ("a", "ADE_m", "ADE / m", "ADE vs Prediction Horizon"),
            ("b", "FDE_m", "FDE / m", "FDE vs Prediction Horizon"),
            ("c", "RMSE_m", "RMSE / m", "RMSE vs Prediction Horizon"),
        ]
        figure_labels = {**MODEL_LABELS, "Helly": "Helly"}
        for ax, (panel, column, ylabel, title) in zip(axes, specs):
            for model in MODEL_ORDER:
                g = time_data[time_data["model_id"].eq(model)].sort_values("rel_time_bin_s")
                ax.plot(
                    g["rel_time_bin_s"], g[column], color=MODEL_COLORS[model],
                    linewidth=2.8 if model == "HGB-ECR" else 2.0,
                    marker="o", markersize=4.5, label=figure_labels[model],
                )
            ax.set_ylabel(ylabel)
            ax.set_title(title, pad=11)
            ax.text(
                0.012, 0.975,
                f"{panel}  {column.replace('_m', '')} over prediction horizon",
                transform=ax.transAxes, ha="left", va="top",
                fontsize=11.0, fontweight="bold",
                bbox={"facecolor": "white", "edgecolor": "none", "alpha": 0.86, "pad": 1.2},
            )
            style_axis(ax, grid_axis="both")
        axes[-1].set_xlabel("Prediction horizon / s")
        handles, labels = axes[0].get_legend_handles_labels()
        fig.legend(handles, labels, loc="upper center", bbox_to_anchor=(0.5, 0.992),
                   ncol=3, columnspacing=1.8, handlelength=2.4,
                   labelspacing=0.65)
        fig.subplots_adjust(
            left=0.10, right=0.98, bottom=0.08, top=0.875, hspace=0.66
        )
    return save_figure(fig, "Figure12_Original_Layout_Larger_Fonts_4K")


def endpoint_data(point):
    p = point.copy()
    p["lateral_error_m"] = (p["pred_x_m"] - p["true_x_m"]).abs()
    p["longitudinal_error_m"] = (p["pred_y_m"] - p["true_y_m"]).abs()
    last = p.sort_values("rel_time_s").groupby(["model_id", "sample_key"], as_index=False).tail(1)
    rows = []
    for model in MODEL_ORDER:
        g = last[last["model_id"].eq(model)]
        lateral = g["lateral_error_m"].mean()
        longitudinal = g["longitudinal_error_m"].mean()
        rows.append({
            "model_id": model,
            "Model": MODEL_LABELS[model],
            "Lateral error (m)": lateral,
            "Longitudinal error (m)": longitudinal,
            "Longitudinal share": longitudinal / (lateral + longitudinal),
            "n": g["sample_key"].nunique(),
        })
    decomposition = pd.DataFrame(rows)
    decomposition.to_csv(SOURCE_DIR / "endpoint_error_decomposition.csv", index=False)
    last.to_csv(SOURCE_DIR / "endpoint_error_by_trajectory.csv", index=False)
    return last, decomposition


def make_figure13(object_metrics, decomposition):
    # Rebuild the original three-panel composition from the archived source
    # data so every label remains legible after Word scales the figure down.
    fig = plt.figure(figsize=(8.0, 6.0))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.18, 1.0], hspace=0.58, wspace=0.50)
    ax = fig.add_subplot(gs[0, :])
    labels = [MODEL_LABELS[model] for model in MODEL_ORDER]
    groups = [
        object_metrics.loc[object_metrics["model_id"].eq(model), "trajectory_end_error_m"]
        .dropna().to_numpy(float)
        for model in MODEL_ORDER
    ]
    bp = ax.boxplot(
        groups, tick_labels=labels, patch_artist=True, showfliers=False, whis=(5, 95),
        medianprops={"color": "#272727", "linewidth": 1.5},
        whiskerprops={"color": "#272727", "linewidth": 1.0},
        capprops={"color": "#272727", "linewidth": 1.0},
        boxprops={"edgecolor": "#4F84B5", "linewidth": 1.0},
    )
    for box in bp["boxes"]:
        box.set_facecolor("#91B8D8")
        box.set_alpha(0.9)
    ax.tick_params(axis="x", rotation=24)
    ax.set_ylabel("FDE / m")
    ax.set_title("Final Displacement Error Distribution", pad=8)
    style_axis(ax)
    panel_label(ax, "a", "FDE distribution")

    ax = fig.add_subplot(gs[1, 0])
    y = np.arange(len(decomposition))
    h = 0.34
    ax.barh(y - h / 2, decomposition["Lateral error (m)"], height=h, color="#BED1E0", label="Lateral")
    ax.barh(y + h / 2, decomposition["Longitudinal error (m)"], height=h, color="#A94442", label="Longitudinal")
    ax.set_yticks(y, decomposition["Model"])
    ax.invert_yaxis()
    ax.set_xlabel("Mean absolute endpoint error (m)")
    handles_b, labels_b = ax.get_legend_handles_labels()
    style_axis(ax, grid_axis="x")
    ax.set_title("b  Endpoint-error decomposition", loc="left", pad=7,
                 fontsize=11.5, fontweight="bold")

    ax = fig.add_subplot(gs[1, 1])
    feature = pd.DataFrame({
        "Feature group": [
            "Longitudinal spacing\nand risk state",
            "Longitudinal\nmotion state",
            "Lateral movement\nstate",
            "Spatial position\nbaseline",
        ],
        "Error increment (m)": [0.239, 0.130, 0.084, 0.081],
        "Contribution (%)": [44.7, 24.4, 15.7, 15.2],
    })
    y = np.arange(len(feature))
    bars = ax.barh(y, feature["Contribution (%)"], color=["#A94442", "#4F789B", "#9CB6CA", "#C4D2DD"])
    ax.set_yticks(y, feature["Feature group"])
    ax.invert_yaxis()
    ax.set_xlabel("Relative contribution (%)")
    for bar, pct, inc in zip(bars, feature["Contribution (%)"], feature["Error increment (m)"]):
        ax.text(pct + 0.8, bar.get_y() + bar.get_height() / 2, f"{pct:.1f}%\n(+{inc:.3f} m)",
                va="center", ha="left", fontsize=10.0)
    ax.set_xlim(0, 62)
    style_axis(ax, grid_axis="x")
    ax.set_title("c  Feature-group contribution", loc="left", pad=7,
                 fontsize=11.5, fontweight="bold")
    fig.legend(
        handles_b, labels_b, loc="center", bbox_to_anchor=(0.27, 0.485),
        ncol=2, fontsize=10.0, handlelength=1.5,
    )
    fig.subplots_adjust(left=0.10, right=0.97, bottom=0.10, top=0.94)
    feature.to_csv(SOURCE_DIR / "hgb_ecr_feature_group_contribution.csv", index=False)
    return save_figure(fig, "Figure13_Endpoint_Feature_4K")


def replace_paragraph_text(paragraph, new_text):
    first_rpr = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            first_rpr = deepcopy(run._r.rPr)
            break
    for child in list(paragraph._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            paragraph._p.remove(child)
    run = paragraph.add_run(new_text)
    if first_rpr is not None:
        is_caption = new_text.startswith(("Figure ", "Table "))
        if paragraph.style.name == "Normal" and not is_caption:
            for tag in ("w:b", "w:bCs", "w:i", "w:iCs"):
                element = first_rpr.find(qn(tag))
                if element is not None:
                    first_rpr.remove(element)
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, first_rpr)


def set_cell_text(cell, text):
    replace_paragraph_text(cell.paragraphs[0], str(text))
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def set_three_line_table(table):
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge, val, size in [
        ("top", "single", "12"), ("left", "nil", "0"), ("bottom", "single", "12"),
        ("right", "nil", "0"), ("insideH", "nil", "0"), ("insideV", "nil", "0"),
    ]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val); element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0"); element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                if row_index < len(table.rows) - 1:
                    paragraph.paragraph_format.keep_with_next = True


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found for restructuring: {prefix}")


def move_elements_after(anchor, elements):
    current = anchor
    for element in elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
        current.addnext(element)
        current = element
    return current


def image_caption_cluster(document, caption_prefix):
    caption = find_paragraph(document, caption_prefix)._p
    body = list(document._element.body)
    caption_index = body.index(caption)
    image_index = None
    for index in range(caption_index - 1, -1, -1):
        if body[index].xpath(".//w:drawing"):
            image_index = index
            break
        if "".join(t.text or "" for t in body[index].xpath(".//w:t")).strip():
            break
    if image_index is None:
        raise RuntimeError(f"Image not found before caption: {caption_prefix}")
    return body[image_index:caption_index + 1]


def table_caption_cluster(document, caption_prefix):
    caption = find_paragraph(document, caption_prefix)._p
    body = list(document._element.body)
    caption_index = body.index(caption)
    table_index = None
    for index in range(caption_index + 1, len(body)):
        if body[index].tag == qn("w:tbl"):
            table_index = index
            break
        text_value = "".join(t.text or "" for t in body[index].xpath(".//w:t")).strip()
        if text_value:
            break
    if table_index is None:
        raise RuntimeError(f"Table not found after caption: {caption_prefix}")
    return body[caption_index:table_index + 1]


def insert_results_subheading(document, before_prefix, title):
    anchor = find_paragraph(document, before_prefix)
    paragraph = anchor.insert_paragraph_before(title, style="Heading 3")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True
        run.italic = False
    return paragraph


def restructure_results_section(document):
    """Interleave evidence with its text and add evidence-led subheadings."""
    # Behavioral indicators: longitudinal evidence, then lateral evidence,
    # followed by scenario-level tests and coupled strategies.
    lateral_anchor = find_paragraph(document, "Figure 9 extends the analysis")._p
    figure9 = image_caption_cluster(document, "Figure 9 Lateral response")
    move_elements_after(lateral_anchor, figure9)

    # Model evaluation: each claim is immediately followed by its figure/table.
    history_anchor = find_paragraph(document, "The history-window sensitivity analysis")._p
    move_elements_after(history_anchor, image_caption_cluster(document, "Figure 10 Representative"))

    performance_anchor = find_paragraph(document, "Figure 11 and Table 9 report")._p
    last = move_elements_after(
        performance_anchor,
        image_caption_cluster(document, "Figure 11 Complete-trajectory"),
    )
    move_elements_after(last, table_caption_cluster(document, "Table 9 Complete-Trajectory"))

    horizon_anchor = find_paragraph(document, "Figure 12 shows the same comparison")._p
    move_elements_after(horizon_anchor, image_caption_cluster(document, "Figure 12 Prediction-error"))

    # Error analysis: keep endpoint evidence together before moving to feature
    # dependence and its separate table.
    endpoint_anchor = find_paragraph(document, "Table 10 confirms that endpoint error")._p
    last = move_elements_after(
        endpoint_anchor,
        image_caption_cluster(document, "Figure 13 Endpoint-error"),
    )
    move_elements_after(last, table_caption_cluster(document, "Table 10 Decomposition"))

    subheadings = [
        ("Figure 8 summarizes the longitudinal motion", "Longitudinal Motion and Risk Indicators"),
        ("Figure 9 extends the analysis", "Lateral Response and Adjacent-Lane Constraints"),
        ("Table 7 reports the scenario medians", "Scenario Differences and Coupled Evasive Strategies"),
        ("The history-window sensitivity analysis", "Model Configuration and Complete-Trajectory Performance"),
        ("Figure 12 shows the same comparison", "Error Evolution over the Prediction Horizon"),
    ]
    for prefix, title in subheadings:
        insert_results_subheading(document, prefix, title)


def fmt_p(p):
    return "<0.001" if p < 0.001 else f"{p:.4f}"


def build_document(figures, detail, record_counts, summary7, tests8, summary_model, decomposition, window, trajectory_change):
    n_records = sum(record_counts.values())
    decel_share = float(np.mean(np.concatenate([
        pd.read_excel(DATA_DIR / f"{day}_断点插值补全.xlsx", sheet_name="Sheet1", usecols=["follower_acc_mps2"])["follower_acc_mps2"].to_numpy(float)
        for day in TRAIN_DAYS
    ]) <= -1.2) * 100)
    decel_events = int(detail["has_decel"].fillna(False).sum())
    corr = float(trajectory_change[["spacing_change_m", "speed_change_mps"]].corr().iloc[0, 1])
    coupled = float(((trajectory_change["spacing_change_m"] < 0) & (trajectory_change["speed_change_mps"] < 0)).mean() * 100)

    no = summary7.iloc[0]
    lead = summary7.iloc[1]
    follow = summary7.iloc[2]
    both = summary7.iloc[3]
    action_pct = pd.crosstab(detail["scene_type"], detail["action"], normalize="index").reindex(
        index=SCENE_ORDER_ZH, columns=ACTION_ORDER_ZH, fill_value=0
    ) * 100
    w1, w3 = window.iloc[0], window.iloc[-1]
    ade_gain = (w1["ADE_m"] - w3["ADE_m"]) / w1["ADE_m"] * 100
    rmse_gain = (w1["RMSE_m"] - w3["RMSE_m"]) / w1["RMSE_m"] * 100
    fde_gain = (w1["FDE_3s_m"] - w3["FDE_3s_m"]) / w1["FDE_3s_m"] * 100

    replacements = {
        "Methods: Field trajectory data were collected during five mobile maintenance operation days in Shanghai with a rear-facing sensing system using radar-camera fusion mounted on the maintenance vehicle. After screening, interpolation, and smoothing, 4,890 valid target trajectories were constructed. Of these, 3,688 training-day trajectories were used for model training, parameter extraction, and behavioral feature analysis, and 1,202 holdout-day validation trajectories were used for formal model evaluation. Variables captured spacing, relative speed, time-to-collision (TTC), deceleration rate to avoid crash (DRAC), lateral displacement, lane position, and adjacent-lane gaps.":
            "Methods: Field trajectory data were collected during five mobile maintenance operation days in Shanghai with a rear-facing radar-camera fusion system mounted on the maintenance vehicle. After screening, interpolation, and smoothing, 4,890 valid target trajectories were constructed. The four training days provided 3,688 trajectories for model development, parameter extraction, and behavioral analysis. The April 9 holdout day contained 1,202 valid trajectories, of which 470 met the complete 3 s history and strict endpoint requirements for matched formal evaluation. Variables captured spacing, relative speed, time-to-collision (TTC), deceleration rate to avoid crash (DRAC), lateral displacement, lane position, and adjacent-lane gaps.",
        "The raw integrated dataset contained 29,310 object IDs. After screening, interpolation, smoothing, and event extraction, the five collection dates yielded 4,890 valid target trajectories: 418 from April 3, 1,725 from April 7, 1,202 from April 9, 1,229 from April 10, and 316 from April 16, 2025. The sample split was date based to avoid evaluating the model on trajectories collected under the same date-specific conditions used for training. Data from April 3, April 7, April 10, and April 16 were used for training, parameter extraction, and behavioral feature analysis, providing 3,688 trajectories. Data from April 9 were reserved as an independent holdout-day validation set, providing 1,202 trajectories for formal model evaluation.":
            "The raw integrated dataset contained 29,310 object IDs. After screening, interpolation, smoothing, and event extraction, the five collection dates yielded 4,890 valid target trajectories: 418 from April 3, 1,725 from April 7, 1,202 from April 9, 1,229 from April 10, and 316 from April 16, 2025. The sample split was date based to avoid evaluating the model under the same date-specific conditions used for training. Data from April 3, April 7, April 10, and April 16 provided 3,688 trajectories for training, parameter extraction, and behavioral analysis. April 9 was retained as the independent holdout day. Of its 1,202 valid trajectories, 470 had a complete 3 s input history and a complete strict evasive endpoint and were used for matched formal evaluation.",
        "Characteristics of Driver Evasive Behavior": "Driver Evasive Behavior Indicator Analysis",
        "This section examines event-level behavioral features extracted from the 3,688 training and feature-analysis trajectories, including onset state, longitudinal speed adjustment, lane-change timing, maneuver duration, and adjacent-lane context. The later model-evaluation results are based on the separate 1,202 holdout-day validation trajectories.":
            "Results follow the variable system and prediction protocol defined in the Methodology. Behavioral indicators were calculated from 3,688 trajectories collected on the four training days, while model performance was evaluated separately on matched trajectories from the April 9 holdout day. The behavioral analysis proceeds from longitudinal motion and risk indicators to lateral response and adjacent-lane constraints, followed by scenario-level differences and coupled evasive strategies. The subsequent model evaluation examines complete-trajectory accuracy, error evolution over the prediction horizon, endpoint-error structure, and predictive feature dependence.",
        "To first describe the overall longitudinal response tendency, the distribution of longitudinal acceleration was examined for the feature-analysis samples. Figure 6 shows that the acceleration distribution was centered below zero, with a mean of -0.67 m/s² and a median of -0.52 m/s². This pattern indicates that longitudinal speed adjustment was common among rear-approaching vehicles during the observed evasive process. The subsequent scenario-level comparison focuses on the event indicators most directly related to approach state, deceleration timing, lane-change timing, maneuver duration, and speed reduction.":
            f"Figure 8 summarizes the longitudinal motion and risk indicators defined in the Methodology. The four training days contained {n_records:,} trajectory records, and {decel_share:.1f}% had longitudinal acceleration at or below the persistent-deceleration threshold of -1.2 m/s². A total of {decel_events:,} evasive events contained an identified persistent-deceleration phase. The joint speed-spacing distribution shows that evasive events began over a broad approach-state range. TTC and deceleration initiation distance also varied across the four moving bottleneck scenarios, indicating that longitudinal response timing depended on the surrounding traffic context.",
        "Figure 6 Distribution of Longitudinal Acceleration":
            "Figure 8 Longitudinal motion and risk indicators in training trajectories",
        "Adjacent-lane constraints were associated with clear differences in both approach state and evasive-response timing. Table 7 summarizes the scenario-level medians for the four target-lane constraint cases. Here, onset spacing and onset speed refer to the spacing and speed at the identified beginning of evasive behavior, not the first frame of the raw trajectory. Vehicles in the following-vehicle-only scenario had the largest onset spacing and the highest onset speed, whereas vehicles in the leading-vehicle-only scenario had the shortest onset spacing and the lowest onset speed. These differences indicate that the target-lane traffic context was already reflected in the vehicle state before the main evasive maneuver began.":
            f"Figure 9 extends the analysis to the lateral and surrounding-vehicle variables. Median lane-change initiation distance increased from {lead['Spacing at lane-change initiation (m)']:.2f} m in the leading-vehicle-only scenario to {both['Spacing at lane-change initiation (m)']:.2f} m under two-sided adjacent-lane constraints. Median lane-change duration was longest in the following-only and two-sided scenarios, at {follow['Lane-change duration (s)']:.2f} s and {both['Lane-change duration (s)']:.2f} s, respectively. Lateral displacement and minimum adjacent-lane gap showed additional dispersion within every scenario, demonstrating that vehicle presence alone did not determine the available evasive space.",
        "The Kruskal-Wallis tests in Table 8 showed significant between-scenario differences for onset spacing, onset speed, TTC, deceleration initiation distance, lane-change initiation distance, lane-change duration, and speed reduction (all p < 0.05). These results support treating adjacent-lane constraints as stratification variables and contextual predictors rather than only as descriptive traffic-context labels.":
            f"Table 7 reports the scenario medians used in Figures 8 and 9. Evasive events began with the largest median spacing under two-sided constraints ({both['Spacing at evasive-event onset (m)']:.2f} m) and the highest median speed when only a following vehicle was present ({follow['Speed at evasive-event onset (m/s)']:.2f} m/s). The following-vehicle-only scenario also produced the largest median speed reduction ({follow['Speed reduction (m/s)']:.2f} m/s). These values show that longitudinal adjustment and lateral timing changed together as adjacent-lane constraints increased.",
        "The median patterns in Table 7 further show how longitudinal and lateral avoidance were coupled. The following-vehicle-only scenario had the longest onset spacing, latest lane-change initiation distance, and largest speed reduction, suggesting that drivers often adjusted speed before entering a target lane constrained from behind. The both-leading-and-following-vehicles scenario had the longest lane-change duration, consistent with the need to coordinate lateral merging within a constrained gap. By contrast, the no-vehicle scenario had the shortest lane-change duration, indicating fewer target-lane conflicts. Figure 7 further illustrates this coupling between spacing change and speed change, showing that evasive behavior cannot be represented only as a longitudinal deceleration process or a lateral lane-changing process.":
            f"The omnibus Kruskal-Wallis tests in Table 8 identified between-scenario differences for all seven indicators (all p < 0.05). The strongest differences occurred in lane-change initiation distance, speed reduction, and spacing at evasive-event onset. Figure 9e further shows that spacing change and speed change were correlated (r = {corr:.2f}); {coupled:.1f}% of trajectories combined decreasing spacing with decreasing speed. Deceleration-then-lane-change was the dominant strategy in every scenario, ranging from {action_pct['减速后换道'].min():.1f}% to {action_pct['减速后换道'].max():.1f}%. Synchronized responses were most frequent under two-sided constraints ({action_pct.loc['相邻侧前后均有车', '同时减速换道']:.1f}%). Pairwise significance was not inferred because no post-hoc test was applied.",
        "Table 7 Behavioral Statistics by Adjacent-Lane Constraint Scenario":
            "Table 7 Behavioral Indicator Statistics under Four Moving Bottleneck Scenarios",
        "Table 8 Kruskal-Wallis Test Results for Behavioral Parameters":
            "Table 8 Differences in Behavioral Indicators among Four Moving Bottleneck Scenarios Based on Kruskal-Wallis Tests",
        "Figure 7 Spacing Change and Speed Change Coupling":
            "Figure 9 Lateral response, surrounding-vehicle constraints, and evasive strategies",
        "These results show that evasive behavior near mobile maintenance vehicles is not a single-lane following problem or a simple lane-changing problem. It is a coupled trajectory process shaped by rear-end risk and adjacent-lane constraints. The following model evaluation therefore tests whether different model families can reproduce this coupled evasive trajectory, rather than only extrapolate short-term motion.":
            "Together, Figures 8 and 9 connect the four variable groups to one behavioral sequence. Drivers adjusted longitudinal risk while evaluating lateral space, and the resulting maneuver depended on both approach state and adjacent-lane constraints. The model evaluation therefore uses complete two-dimensional evasive trajectories rather than isolated braking or lane-changing outcomes.",
        "Figure 8 first illustrates the trajectory-level differences among the compared models. In the representative lane-changing and continuous lateral-adjustment samples, Kalman-CV and Helly Baseline mainly extended the recent motion state and therefore did not sufficiently capture the lateral avoidance process. IDM-Avoidance represented part of the avoidance logic but remained sensitive to rule thresholds, while GBR and 1D-CNN showed larger deviations as prediction errors accumulated. By contrast, HGB-ECR stayed closer to the observed evasive trajectories because the constant-velocity anchor provided a stable reference, the residual module corrected nonlinear deviations, and the endpoint correction module reduced terminal displacement error.":
            f"The history-window sensitivity analysis supported the 3 s setting used by all models. Relative to the 1 s setting, the 3 s window reduced ADE by {ade_gain:.1f}%, RMSE by {rmse_gain:.1f}%, and 3 s FDE by {fde_gain:.1f}% on the common matched holdout subset. Figure 10 then shows the original representative lane-changing and continuous lateral-adjustment cases. HGB-ECR remained closest to the observed trajectories during the later maneuver stages, whereas the baselines showed larger longitudinal or lateral deviations.",
        "The overall quantitative results are reported in Table 9 and visualized in Figure 9. The metrics in Table 9 were calculated over the predicted evasive behavior trajectory, rather than at a single time point. HGB-ECR achieved the lowest overall error on the 1,202 holdout-day validation trajectories, with an ADE of 3.854 m, an FDE of 5.933 m, and an RMSE of 5.863 m. Compared with the strongest baseline for each metric, HGB-ECR reduced ADE by 44.7%, FDE by 39.9%, and RMSE by 43.0%. The grouped comparison in Figure 9 further shows that HGB-ECR consistently produced the smallest ADE, FDE, and RMSE, whereas Helly Baseline had the largest overall prediction error. These results indicate that HGB-ECR improved both whole-trajectory fit and terminal-position representation for the evaluated evasive trajectories.":
            "Figure 11 and Table 9 report complete-trajectory performance on the 470 matched holdout trajectories and 5,295 evaluation points. HGB-ECR achieved the lowest error for every metric, with an ADE of 3.854 m, an FDE of 5.933 m, and an RMSE of 5.863 m. Kalman-CV was the strongest baseline for ADE and RMSE, whereas GBR was the strongest baseline for FDE. Relative to these metric-specific baselines, HGB-ECR reduced ADE by 44.7%, FDE by 39.9%, and RMSE by 43.0%. Helly Baseline produced the largest errors, confirming the limitation of a primarily longitudinal response model for coupled evasive trajectories.",
        "Figure 10 further shows how prediction error evolved with the prediction horizon. Although all models showed increasing ADE, FDE, and RMSE as the horizon became longer, HGB-ECR maintained the lowest error curves as the rolling prediction horizon increased. This pattern indicates that its advantage was not limited to the aggregate values in Table 9. Instead, HGB-ECR also controlled error accumulation more effectively as the evasive maneuver unfolded, especially during the later prediction horizons when deceleration timing, lateral adjustment, and endpoint position became more uncertain.":
            "Figure 12 shows the same comparison over the rolling prediction horizon. Errors increased with time for every model as later predictions depended on previously predicted states. HGB-ECR maintained the lowest ADE, FDE, and RMSE curves across the evaluated horizons. Its separation from the baselines became clearer during the later steps, when uncertainty in deceleration timing, forward progression, and lateral-state transitions accumulated.",
        "Figure 8 Representative Observed and Modeled Evasive Trajectories":
            "Figure 10 Representative observed and modeled evasive trajectories",
        "Figure 9 Overall Trajectory Prediction Error across Models":
            "Figure 11 Complete-trajectory prediction performance across models",
        "Table 9 Model Evaluation Performance on the Holdout-Day Validation Set":
            "Table 9 Complete-Trajectory Prediction Performance on the Holdout-Day Validation Set",
        "Figure 10 Time-Horizon Evolution of Prediction Error for Driver Evasive Behavior":
            "Figure 12 Prediction-error evolution over the rolling horizon",
        "The preceding evaluation showed that HGB-ECR achieved the lowest prediction errors. Figure 11 further shows that its endpoint errors were more stable across samples. HGB-ECR had the lowest median final displacement error and a relatively compact distribution, whereas Helly Baseline showed the largest error spread. This indicates that HGB-ECR improved both endpoint accuracy and prediction stability.":
            "Figure 13a compares endpoint-error distributions across the six models. HGB-ECR combined the lowest median FDE with a relatively compact distribution across the matched trajectories. Figure 13b shows that its advantage was concentrated in the longitudinal component, while lateral endpoint errors remained small for several models.",
        "Table 10 decomposes the endpoint error into lateral and longitudinal components. For all models, the endpoint error was mainly longitudinal. For HGB-ECR, the lateral error was 0.851 m and the longitudinal error was 5.702 m, with the longitudinal component accounting for 87.0% of the endpoint error. This result shows that the main difficulty was predicting forward progression, deceleration timing, and final closing distance, rather than lateral lane position alone.":
            "Table 10 confirms that endpoint error was predominantly longitudinal for every model, accounting for 85.4% to 96.0% of the component sum. HGB-ECR had a lateral error of 0.851 m and the lowest longitudinal error, at 5.702 m. It did not have the smallest lateral error, but it reduced longitudinal error substantially relative to every baseline. The principal modeling challenge was therefore forward progression and deceleration timing rather than lane position alone.",
        "Table 11 reports the feature-group perturbation results for HGB-ECR. Perturbing longitudinal spacing and risk state produced the largest increase in endpoint prediction error, 0.239 m, accounting for 44.7% of the total contribution. Longitudinal motion state ranked second at 24.4%, followed by lateral movement state and the spatial position baseline. These results indicate that longitudinal risk and motion information contributed most to HGB-ECR prediction performance in this dataset.":
            "The feature-group perturbation results in Figure 13c and Table 11 follow this error structure. Perturbing longitudinal spacing and risk variables increased endpoint error by 0.239 m and accounted for 44.7% of total positive contribution. Longitudinal motion variables ranked second at 24.4%. Lateral movement state and the spatial position baseline contributed 15.7% and 15.2%, respectively. HGB-ECR therefore relied most strongly on variables describing approach pressure and longitudinal state.",
        "The perturbation results should be interpreted as predictive evidence, not causal attribution. They show which variables were most useful for model prediction in this dataset, but they do not prove that these variables independently caused driver decisions.":
            "The perturbation analysis measures predictive dependence rather than causal influence. It identifies which feature groups supported HGB-ECR performance in this dataset, but it does not establish that any variable independently caused a driver decision. Together, Figure 13 and Tables 10 and 11 locate the remaining uncertainty mainly in the longitudinal evolution of evasive events.",
        "Figure 11 Endpoint Error Distribution across Models":
            "Figure 13 Endpoint-error structure and HGB-ECR feature dependence",
        "Table 10 Lateral and Longitudinal Error Decomposition":
            "Table 10 Decomposition of Endpoint Error into Lateral and Longitudinal Components",
        "Table 11 Feature-Group Contribution for HGB-ECR":
            "Table 11 Feature-Group Perturbation Results for HGB-ECR",
        "The results provide a basis for short-horizon risk identification behind mobile maintenance vehicles. Rear-approaching vehicle risk should not be judged only by current spacing or by TTC and DRAC alone. It also depends on whether the vehicle is still closing in, whether deceleration is sufficient, whether lateral avoidance has started, and whether the adjacent lane provides enough space for merging.":
            "The results support a multi-indicator approach to short-horizon risk identification behind an MMV. Current spacing, TTC, and DRAC describe longitudinal pressure, but they do not show whether deceleration is sufficient or lateral avoidance has begun. Adjacent-lane gaps and the predicted two-dimensional trajectory provide the additional context needed to distinguish continued approach, controlled deceleration, and constrained lane changing.",
        "HGB-ECR can support this assessment by predicting the short-horizon evasive trajectory. The predicted trajectory can help identify three types of vehicles: vehicles that continue approaching the maintenance vehicle without sufficient deceleration, vehicles that show lateral avoidance but face limited adjacent-lane gaps, and vehicles whose predicted paths remain uncertain because of sudden braking or unstable lateral movement. These cases are directly relevant to rear-end warning and lane-change conflict identification.":
            "HGB-ECR can support this assessment by predicting the short-horizon evasive path. The output can flag vehicles that continue closing without sufficient deceleration, vehicles that begin lateral avoidance under limited gaps, and vehicles with unstable paths after sudden braking or strong lateral movement. These categories correspond to rear-end warning, lane-change conflict identification, and conservative handling of uncertain predictions.",
        "The model should be used as a decision-support tool rather than as an independent safety rule. In field deployment, predicted trajectories should be combined with real-time sensing quality checks, TTC, DRAC, adjacent-lane gap information, conservative warning thresholds, and human safety judgment. Further validation is still needed across different roadways, traffic densities, weather conditions, and maintenance operation types.":
            "The model remains a decision-support tool rather than an independent safety rule. Deployment should combine predicted trajectories with sensing-quality checks, TTC, DRAC, adjacent-lane gaps, conservative warning thresholds, and human oversight. The present formal evaluation contains 470 matched trajectories from one holdout day. Cross-roadway, cross-operation, traffic-density, weather, and lighting validation is required before broader use.",
        "This study developed a field-trajectory-based framework to model driver evasive behavior under mobile maintenance operations on expressways. Field trajectory data were collected during five mobile maintenance operation days in Shanghai using a rear-facing radar-camera fusion sensing system mounted on the maintenance vehicle. After screening, interpolation, and smoothing, 4,890 valid target trajectories were constructed. Among them, 3,688 training-day trajectories were used for model training, parameter extraction, and behavioral feature analysis, and 1,202 holdout-day validation trajectories were used for formal model evaluation. The study extracted evasive-behavior features, compared behavioral differences under adjacent-lane constraints, and evaluated HGB-ECR against kinematic, car-following, rule-based, and data-driven baselines. Several key conclusions can be drawn:":
            "This study developed a field-trajectory-based framework to model driver evasive behavior under mobile maintenance operations on expressways. Five mobile maintenance operation days produced 4,890 valid target trajectories after screening, interpolation, and smoothing. The four training days provided 3,688 trajectories for model development and behavioral analysis. The April 9 holdout day contained 1,202 valid trajectories, of which 470 met the complete-history and strict-endpoint requirements for matched formal evaluation. The study compared behavioral indicators under adjacent-lane constraints and evaluated HGB-ECR against kinematic, car-following, rule-based, and data-driven baselines. Several key conclusions can be drawn:",
        "HGB-ECR achieved the lowest trajectory prediction errors on the 1,202 holdout-day validation trajectories. Its ADE, FDE, and RMSE were 3.854 m, 5.933 m, and 5.863 m, respectively. Compared with the strongest baseline for each metric, HGB-ECR reduced ADE by 44.7%, FDE by 39.9%, and RMSE by 43.0%. These results suggest that constant-velocity anchoring, residual histogram gradient boosting, and endpoint correction improve short-horizon prediction of evasive trajectories.":
            "HGB-ECR achieved the lowest complete-trajectory prediction errors on the 470 matched holdout trajectories. Its ADE, FDE, and RMSE were 3.854 m, 5.933 m, and 5.863 m, respectively. Compared with the strongest baseline for each metric, HGB-ECR reduced ADE by 44.7%, FDE by 39.9%, and RMSE by 43.0%. These results suggest that constant-velocity anchoring, residual histogram gradient boosting, and endpoint correction improve short-horizon prediction of evasive trajectories.",
    }

    document = Document(SOURCE_DOCX)
    found = set()
    for paragraph in document.paragraphs:
        key = paragraph.text.strip()
        is_figure_caption = key in {
            "Figure 6 Distribution of Longitudinal Acceleration",
            "Figure 7 Spacing Change and Speed Change Coupling",
            "Figure 8 Representative Observed and Modeled Evasive Trajectories",
            "Figure 9 Overall Trajectory Prediction Error across Models",
            "Figure 10 Time-Horizon Evolution of Prediction Error for Driver Evasive Behavior",
            "Figure 11 Endpoint Error Distribution across Models",
        }
        is_table_caption = key in {
            "Table 7 Behavioral Statistics by Adjacent-Lane Constraint Scenario",
            "Table 8 Kruskal-Wallis Test Results for Behavioral Parameters",
            "Table 9 Model Evaluation Performance on the Holdout-Day Validation Set",
            "Table 10 Lateral and Longitudinal Error Decomposition",
            "Table 11 Feature-Group Contribution for HGB-ECR",
        }
        if key in replacements:
            replace_paragraph_text(paragraph, replacements[key])
            found.add(key)
            if is_figure_caption:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
            elif not is_table_caption and len(replacements[key]) > 180:
                # Prevent prose paragraphs that begin with "Figure" or
                # "Table" from inheriting caption boldface from the source.
                for run in paragraph.runs:
                    run.bold = False
                    run.italic = False
        if paragraph.text.startswith("Table "):
            paragraph.paragraph_format.keep_with_next = True
    missing = sorted(set(replacements) - found)
    if missing:
        raise RuntimeError("Expected source paragraphs were not found:\n" + "\n".join(missing))

    # Table 7 (document.tables[4])
    table7 = document.tables[4]
    headers7 = list(summary7.columns)
    for cell, value in zip(table7.rows[0].cells, headers7):
        set_cell_text(cell, value)
    for row, (_, values) in zip(table7.rows[1:], summary7.iterrows()):
        for cell, column in zip(row.cells, headers7):
            value = values[column]
            if column == "Evasive events, n":
                text = f"{int(value)}"
            elif column == "Moving bottleneck scenario":
                text = str(value)
            else:
                text = f"{float(value):.2f}"
            set_cell_text(cell, text)

    # Table 8 (document.tables[5])
    table8 = document.tables[5]
    headers8 = ["Behavioral indicator", "H statistic", "p value", "Decision at α = 0.05"]
    for cell, value in zip(table8.rows[0].cells, headers8):
        set_cell_text(cell, value)
    for row, (_, values) in zip(table8.rows[1:], tests8.iterrows()):
        set_cell_text(row.cells[0], values["Behavioral indicator"])
        set_cell_text(row.cells[1], f"{values['H statistic']:.2f}")
        set_cell_text(row.cells[2], fmt_p(values["p value"]))
        set_cell_text(row.cells[3], "Reject H₀")

    # Table 9 (document.tables[6])
    table9 = document.tables[6]
    set_cell_text(table9.rows[0].cells[5], "n")
    summary_lookup = summary_model.set_index("model_id")
    table_model_ids = ["HGB-ECR", "Kalman-CV", "GBR", "IDM", "CNN", "Helly"]
    for row, model in zip(table9.rows[1:], table_model_ids):
        values = summary_lookup.loc[model]
        set_cell_text(row.cells[0], MODEL_LABELS[model])
        set_cell_text(row.cells[2], f"{values['ADE_m']:.3f}")
        set_cell_text(row.cells[3], f"{values['FDE_m']:.3f}")
        set_cell_text(row.cells[4], f"{values['RMSE_m']:.3f}")
        set_cell_text(row.cells[5], f"{int(values['N'])}")

    # Table 10 (document.tables[7])
    table10 = document.tables[7]
    set_cell_text(table10.rows[0].cells[4], "n")
    decomp_lookup = decomposition.set_index("model_id")
    for row, model in zip(table10.rows[1:], table_model_ids):
        values = decomp_lookup.loc[model]
        set_cell_text(row.cells[0], MODEL_LABELS[model])
        set_cell_text(row.cells[1], f"{values['Lateral error (m)']:.3f}")
        set_cell_text(row.cells[2], f"{values['Longitudinal error (m)']:.3f}")
        set_cell_text(row.cells[3], f"{values['Longitudinal share'] * 100:.1f}%")
        set_cell_text(row.cells[4], f"{int(values['n'])}")

    for table_index in range(4, 9):
        set_three_line_table(document.tables[table_index])

    # Resize Results figures to match the new panel aspect ratios.
    target_width = min(
        document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin,
        Inches(7.0),
    )
    ratios = [8.0 / 5.6, 8.0 / 7.8, 4096 / 1954, 4857 / 2817, 3847 / 4515, 8.0 / 6.0]
    max_height = Inches(7.0)
    for shape, ratio in zip(list(document.inline_shapes)[7:13], ratios):
        width = min(target_width, int(max_height * ratio))
        shape.width = width
        shape.height = int(width / ratio)

    document.save(TEMP_DOCX)

    media_replacements = {
        "word/media/image48.png": figures[0],
        "word/media/image49.png": figures[1],
        "word/media/image50.png": figures[2],
        "word/media/image51.png": figures[3],
        "word/media/image52.png": figures[4],
        "word/media/image53.png": figures[5],
    }
    with ZipFile(TEMP_DOCX, "r") as zin, ZipFile(TEMP_MEDIA_DOCX, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in media_replacements:
                data = Path(media_replacements[item.filename]).read_bytes()
            zout.writestr(item, data)
    os.replace(TEMP_MEDIA_DOCX, OUTPUT_DOCX)
    TEMP_DOCX.unlink(missing_ok=True)

    # The requested deliverable contains the Results section only. Remove all
    # content before RESULTS AND DISCUSSION and from CONCLUSIONS onward while
    # preserving the section properties and embedded figure media.
    results_doc = Document(OUTPUT_DOCX)
    body = results_doc._element.body
    children = list(body)
    start = end = None
    for index, child in enumerate(children):
        texts = child.xpath(".//w:t")
        text_value = "".join(t.text or "" for t in texts).strip()
        if text_value == "RESULTS AND DISCUSSION" and start is None:
            start = index
        elif text_value == "CONCLUSIONS" and start is not None:
            end = index
            break
    if start is None or end is None:
        raise RuntimeError("Could not isolate the Results section boundaries")
    for index, child in reversed(list(enumerate(children))):
        if child.tag == qn("w:sectPr"):
            continue
        if index < start or index >= end:
            body.remove(child)
    restructure_results_section(results_doc)
    results_doc.save(OUTPUT_DOCX)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--figures-only", action="store_true",
        help="Generate the six 4K Results figures without editing a DOCX.",
    )
    args = parser.parse_args()
    apply_style()
    accelerations, trajectory_change, detail, record_counts = read_training_data()
    if args.figures_only:
        figure = make_figure9(detail, trajectory_change)
        print(figure)
        return
    summary7, tests8 = scene_summary_and_tests(detail)
    summary_model, time_data, point, object_metrics, window = load_model_results()
    last, decomposition = endpoint_data(point)

    figures = [
        make_figure8(accelerations, detail, record_counts),
        make_figure9(detail, trajectory_change),
        make_figure10(point, window),
        make_figure11(summary_model),
        make_figure12(time_data),
        make_figure13(object_metrics, decomposition),
    ]
    build_document(
        figures, detail, record_counts, summary7, tests8,
        summary_model, decomposition, window, trajectory_change,
    )
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
